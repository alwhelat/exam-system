"""
=============================================================
نظام توزيع مراقبي الامتحانات النهائية
Exam Monitor Distribution System - v2.0
=============================================================
التشغيل / Run:
    python exam_monitor_system.py
ثم افتح / Then open:
    http://localhost:5000
=============================================================
"""

import os
import json
import re
from pathlib import Path
from datetime import datetime

import pandas as pd
from flask import Flask, render_template_string, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

app = Flask(__name__)

# ─────────────────────────────────────────────
# FILE PATHS
# ─────────────────────────────────────────────
BASE_DIR   = Path(r"D:\FUC\اللجنة الامتحانية\2025-2026\منصة مراقبات")
STAFF_FILE = BASE_DIR / "التدريسين - 2025-2026.xlsx"
HALLS_FILE = BASE_DIR / "اسماء القاعات.xlsx"
HTML_FILE  = BASE_DIR / "نظام_المراقبين_v2.html"

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
DAYS_ALL = ['السبت', 'الاحد', 'الاثنين', 'الثلاثاء', 'الاربعاء', 'الخميس']

TITLE_LABEL = {
    'prof':     'أ.د',
    'asst_prof': 'أ.م.د',
    'lec_dr':   'م.د',
    'dr':       'د.',
    'asst':     'أ.م',
    'mm':       'م.م',
    'other':    '—'
}

POS_RANK = {
    'تدريسي': 3, 'تدريسية': 3,
    'مقرر قسم': 2, 'رئيس قسم': 1,
    'مقرر كلية': 0, 'عميد': 0, 'معاون عميد': 0
}

FEM_NAMES = {
    'هند', 'ليلى', 'منى', 'هبة', 'مريم', 'سيناء', 'ايالف', 'رنا', 'لميس', 'رونق', 'ضحى',
    'رغد', 'ياسمين', 'استبرق', 'نور', 'فاطمة', 'امل', 'نبأ', 'بان', 'هيفاء', 'ريما',
    'شهد', 'نغم', 'رحمه', 'رغدة', 'سارة', 'مينا', 'همسة', 'سجى', 'ذكريات', 'تبارك',
    'علياء', 'ايثار', 'وداد', 'هدى', 'فرح', 'غفران', 'زينب', 'ميسرة', 'روى', 'لبنى',
    'ايلاف', 'زهراء', 'بنين', 'تقوى', 'اشواق', 'مروة', 'اية', 'شذى', 'الاء', 'نورا',
    'هالة', 'ريا', 'الهام', 'ورود', 'انوار', 'هناء', 'سلوى', 'كفاء', 'لقاء', 'دينا',
    'زينة', 'رند', 'شيماء', 'منال', 'نورالهدى', 'فهيمة', 'اتحاد', 'ماجدة', 'سرى', 'سحر',
    'صفاء', 'ايمان', 'روضة', 'حنان', 'خلود', 'سماء', 'عبير', 'رهام', 'وجدان', 'نادية'
}

# ─────────────────────────────────────────────
# DATA PROCESSING
# ─────────────────────────────────────────────

def get_title_info(name: str):
    n = name.strip()
    checks = [
        ('أ.د',   'prof',     6, True),
        ('ا.د',   'prof',     6, True),
        ('أ.م.د', 'asst_prof', 5, True),
        ('ا.م.د', 'asst_prof', 5, True),
        ('م.د',   'lec_dr',   4, True),
        ('د.',    'dr',       4, True),
        ('أ.م',   'asst',     3, False),
        ('ا.م',   'asst',     3, False),
        ('م.م.',  'mm',       2, False),
        ('م.م',   'mm',       2, False),
        ('م. ',   'mm',       2, False),
        ('م.',    'mm',       2, False),
    ]
    for pfx, code, rank, phd in checks:
        if n.startswith(pfx):
            return code, rank, phd
    return 'other', 1, False


def guess_gender(name: str) -> str:
    parts = name.split()
    for p in parts[1:4]:
        clean = re.sub(r'[.ال]', '', p).strip()
        if clean in FEM_NAMES or p in FEM_NAMES:
            return 'female'
    return 'male'


def parse_days(day_str) -> list:
    if pd.isna(day_str):
        return []
    s = str(day_str).strip()
    if s == 'كامل':
        return DAYS_ALL[:]
    return [d.strip() for d in re.split(r'[,،]', s) if d.strip()]


def load_staff(filepath: Path) -> list:
    df = pd.read_excel(filepath, sheet_name='الملاك')
    records = []
    for i, row in df.iterrows():
        name = str(row['اسم التدريسي']).strip()
        tc, tr, phd = get_title_info(name)
        mtype = str(row.get('نوع الدوام', 'محاضر')).strip()
        if pd.isna(row.get('نوع الدوام', '')):
            mtype = 'محاضر'
        pos = str(row.get('المنصب', 'تدريسي')).strip()
        if pd.isna(row.get('المنصب', '')):
            pos = 'تدريسي'
        days = parse_days(row.get('ايام الدوام الاسبوعي', ''))
        in_comm = str(row.get('هل لجنة امتحانية؟', 'لا')).strip() == 'نعم'
        gender = guess_gender(name)
        pr = POS_RANK.get(pos, 1)
        typr = 10 if mtype == 'ملاك' else 0
        records.append({
            'id': i,
            'name': name,
            'college': str(row.get('الكلية', '')).strip(),
            'dept': str(row.get('القسم', '')).strip(),
            'title_code': tc,
            'title_rank': tr,
            'has_phd': phd,
            'mtype': mtype,
            'type_rank': typr,
            'pos': pos,
            'pos_rank': pr,
            'days': days,
            'in_committee': in_comm,
            'gender': gender,
            'phone': str(row.get('رقم الهاتف', '')).strip(),
            'total_rank': typr + pr + tr,
            'is_reserve': False,
            'is_manual': False,
        })
    return records


def load_halls(filepath: Path) -> list:
    df = pd.read_excel(filepath)
    return [
        {
            'name': str(row['اسم القاعة']).strip(),
            'monitors_needed': int(row['العدد المطلوب من المراقبين']),
            'is_extra': False
        }
        for _, row in df.iterrows()
    ]


# ─────────────────────────────────────────────
# DISTRIBUTION ENGINE
# ─────────────────────────────────────────────

class DistributionEngine:

    def __init__(self, staff: list, halls: list):
        self.staff = staff
        self.halls = halls

    def get_available(self, day: str, exclude_committee: bool = True,
                      exclude_reserve: bool = False) -> list:
        result = []
        for s in self.staff:
            if exclude_committee and s['in_committee']:
                continue
            if exclude_reserve and s['is_reserve']:
                continue
            if day in s['days']:
                result.append(s)
        return sorted(result, key=lambda x: -x['total_rank'])

    def run(self, day: str, selected_halls: list,
            exclude_committee: bool = True,
            exclude_reserve: bool = False) -> dict:
        avail = self.get_available(day, exclude_committee, exclude_reserve)
        supervisors = [s for s in avail if s['has_phd']]
        monitors = [s for s in avail if s['title_code'] in ('mm', 'asst')]
        males = [s for s in monitors if s['gender'] == 'male']
        females = [s for s in monitors if s['gender'] == 'female']

        used_sup: set = set()
        used_mon: set = set()
        distribution = {}

        hall_list = [h for h in self.halls if h['name'] in selected_halls]

        for hall in hall_list:
            needed = hall['monitors_needed']
            sup = next((s for s in supervisors
                        if s['id'] not in used_sup and s['id'] not in used_mon), None)
            if sup:
                used_sup.add(sup['id'])

            assigns = []
            slots = needed - 1

            if slots >= 2:
                m = next((s for s in males
                          if s['id'] not in used_mon and s['id'] not in used_sup), None)
                f = next((s for s in females
                          if s['id'] not in used_mon and s['id'] not in used_sup
                          and s['id'] != (m['id'] if m else -1)), None)
                if m:
                    assigns.append(m)
                    used_mon.add(m['id'])
                if f:
                    assigns.append(f)
                    used_mon.add(f['id'])
                for _ in range(len(assigns), slots):
                    nx = next((s for s in monitors
                               if s['id'] not in used_mon and s['id'] not in used_sup), None)
                    if nx:
                        assigns.append(nx)
                        used_mon.add(nx['id'])
                    else:
                        assigns.append(None)
            elif slots == 1:
                nx = next((s for s in monitors
                           if s['id'] not in used_mon and s['id'] not in used_sup), None)
                if nx:
                    assigns.append(nx)
                    used_mon.add(nx['id'])
                else:
                    assigns.append(None)

            distribution[hall['name']] = {
                'supervisor': sup,
                'monitors': assigns,
                'hall': hall
            }

        return distribution

    def get_remaining(self, day: str, distribution: dict,
                      exclude_committee: bool = True) -> list:
        assigned_ids = set()
        for d in distribution.values():
            if d['supervisor']:
                assigned_ids.add(d['supervisor']['id'])
            for m in d['monitors']:
                if m:
                    assigned_ids.add(m['id'])

        return [
            s for s in self.staff
            if day in s['days']
            and s['id'] not in assigned_ids
            and not (exclude_committee and s['in_committee'])
        ]


# ─────────────────────────────────────────────
# WORD DOCUMENT GENERATOR
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def add_bordered_table(doc, rows_data, header_row, col_widths):
    table = doc.add_table(rows=0, cols=len(header_row))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.add_row()
    for i, (cell, w) in enumerate(zip(hdr.cells, col_widths)):
        cell.width = Cm(w)
        cell.text = header_row[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1E3A8A')

    for ri, row_data in enumerate(rows_data):
        row = table.add_row()
        bg = 'EEF2FF' if ri % 2 == 0 else 'FFFFFF'
        for i, (cell, val) in enumerate(zip(row.cells, row_data)):
            cell.width = Cm(col_widths[i])
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            set_cell_bg(cell, bg)

    return table


def generate_word_document(distribution: dict, remaining: list,
                           day: str, date: str, session: str) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('جدول توزيع مراقبي الامتحانات النهائية')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x6E)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"اليوم: {day}    التاريخ: {date or '___________'}"
        + (f"    الدورة: {session}" if session else '')
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    headers = ['القاعة', 'المطلوب', 'الدور', 'الاسم', 'النوع', 'اللقب', 'الجنس']
    col_w = [2.5, 1.5, 2.5, 5.5, 1.8, 1.8, 1.8]
    rows = []
    for hall_name, d in distribution.items():
        sup = d['supervisor']
        hall = d['hall']
        rows.append([
            hall_name,
            hall['monitors_needed'],
            'مشرف القاعة',
            sup['name'] if sup else '—',
            sup['mtype'] if sup else '—',
            TITLE_LABEL.get(sup['title_code'], '—') if sup else '—',
            'ذكر' if sup and sup['gender'] == 'male' else ('أنثى' if sup else '—')
        ])
        for i, mon in enumerate(d['monitors']):
            rows.append([
                '',
                '',
                f'مراقب {i+1}',
                mon['name'] if mon else '—',
                mon['mtype'] if mon else '—',
                TITLE_LABEL.get(mon['title_code'], '—') if mon else '—',
                'ذكر' if mon and mon['gender'] == 'male' else ('أنثى' if mon else '—')
            ])

    add_bordered_table(doc, rows, headers, col_w)
    doc.add_paragraph()

    if remaining:
        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = h2.add_run(f'المراقبون غير الموزَّعين اليوم ({len(remaining)} شخص)')
        run2.font.size = Pt(13)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

        rem_headers = ['#', 'الاسم', 'الكلية', 'النوع', 'اللقب', 'الجنس', 'الهاتف']
        rem_w = [0.8, 5.5, 3.5, 1.8, 1.5, 1.5, 3.0]
        rem_rows = [
            [i + 1, s['name'], s['college'], s['mtype'],
             TITLE_LABEL.get(s['title_code'], '—'),
             'ذكر' if s['gender'] == 'male' else 'أنثى',
             s.get('phone', '—') or '—']
            for i, s in enumerate(remaining)
        ]
        add_bordered_table(doc, rem_rows, rem_headers, rem_w)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(
        f"تم الإنشاء: {datetime.now().strftime('%Y/%m/%d %H:%M')}  |  "
        f"إجمالي القاعات: {len(distribution)}  |  "
        f"إجمالي المتبقين: {len(remaining)}"
    )
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# LOAD DATA
# ─────────────────────────────────────────────
STAFF_DATA = []
HALLS_DATA = []


def load_data():
    global STAFF_DATA, HALLS_DATA
    if STAFF_FILE.exists():
        STAFF_DATA = load_staff(STAFF_FILE)
        print(f"✓ تم تحميل {len(STAFF_DATA)} تدريسي من: {STAFF_FILE.name}")
    else:
        print(f"✗ ملف التدريسيين غير موجود: {STAFF_FILE}")

    if HALLS_FILE.exists():
        HALLS_DATA = load_halls(HALLS_FILE)
        print(f"✓ تم تحميل {len(HALLS_DATA)} قاعة من: {HALLS_FILE.name}")
    else:
        print(f"✗ ملف القاعات غير موجود: {HALLS_FILE}")


# ─────────────────────────────────────────────
# FLASK ROUTES
# ─────────────────────────────────────────────

@app.route('/')
def index():
    if HTML_FILE.exists():
        return HTML_FILE.read_text(encoding='utf-8')
    return "<h1>خطأ: ملف الواجهة غير موجود</h1><p>" + str(HTML_FILE) + "</p>", 404


@app.route('/api/staff', methods=['GET'])
def api_get_staff():
    return jsonify(STAFF_DATA)


@app.route('/api/staff', methods=['POST'])
def api_add_staff():
    data = request.json
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': 'الاسم مطلوب'}), 400
    tc, tr, phd = get_title_info(name)
    mtype = data.get('mtype', 'ملاك')
    pos = data.get('pos', 'تدريسي')
    pr = POS_RANK.get(pos, 1)
    typr = 10 if mtype == 'ملاك' else 0
    new_id = max((s['id'] for s in STAFF_DATA), default=0) + 1
    new_s = {
        'id': new_id,
        'name': name,
        'college': data.get('college', ''),
        'dept': data.get('dept', ''),
        'title_code': tc, 'title_rank': tr, 'has_phd': phd,
        'mtype': mtype, 'type_rank': typr,
        'pos': pos, 'pos_rank': pr,
        'days': data.get('days', DAYS_ALL),
        'in_committee': data.get('in_committee', False),
        'gender': guess_gender(name),
        'phone': data.get('phone', ''),
        'total_rank': typr + pr + tr,
        'is_reserve': data.get('is_reserve', False),
        'is_manual': True,
    }
    STAFF_DATA.append(new_s)
    return jsonify(new_s), 201


@app.route('/api/staff/<int:sid>', methods=['DELETE'])
def api_delete_staff(sid):
    global STAFF_DATA
    STAFF_DATA = [s for s in STAFF_DATA if s['id'] != sid]
    return jsonify({'ok': True})


@app.route('/api/staff/<int:sid>/reserve', methods=['PATCH'])
def api_toggle_reserve(sid):
    s = next((x for x in STAFF_DATA if x['id'] == sid), None)
    if not s:
        return jsonify({'error': 'not found'}), 404
    s['is_reserve'] = not s.get('is_reserve', False)
    return jsonify(s)


@app.route('/api/halls', methods=['GET'])
def api_get_halls():
    return jsonify(HALLS_DATA)


@app.route('/api/halls', methods=['POST'])
def api_add_hall():
    data = request.json
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': 'اسم القاعة مطلوب'}), 400
    if any(h['name'] == name for h in HALLS_DATA):
        return jsonify({'error': 'القاعة موجودة مسبقاً'}), 400
    new_h = {
        'name': name,
        'monitors_needed': int(data.get('monitors_needed', 2)),
        'is_extra': True
    }
    HALLS_DATA.append(new_h)
    return jsonify(new_h), 201


@app.route('/api/halls/<name>', methods=['DELETE'])
def api_delete_hall(name):
    global HALLS_DATA
    HALLS_DATA = [h for h in HALLS_DATA if h['name'] != name]
    return jsonify({'ok': True})


@app.route('/api/distribute', methods=['POST'])
def api_distribute():
    data = request.json
    day = data.get('day', 'السبت')
    selected = data.get('selected_halls', [h['name'] for h in HALLS_DATA])
    excm = data.get('exclude_committee', True)
    exres = data.get('exclude_reserve', False)
    manual_overrides = data.get('manual_overrides', {})

    engine = DistributionEngine(STAFF_DATA, HALLS_DATA)
    dist = engine.run(day, selected, excm, exres)

    staff_map = {s['id']: s for s in STAFF_DATA}
    for hall_name, override in manual_overrides.items():
        if hall_name not in dist:
            continue
        if 'sup_id' in override:
            dist[hall_name]['supervisor'] = staff_map.get(override['sup_id'])
        if 'mon_ids' in override:
            dist[hall_name]['monitors'] = [staff_map.get(mid) for mid in override['mon_ids']]

    remaining = engine.get_remaining(day, dist, excm)

    def ser(person):
        if not person:
            return None
        return {k: v for k, v in person.items()}

    result = {}
    for k, v in dist.items():
        result[k] = {
            'supervisor': ser(v['supervisor']),
            'monitors': [ser(m) for m in v['monitors']],
            'hall': v['hall']
        }

    return jsonify({
        'distribution': result,
        'remaining': [ser(s) for s in remaining],
        'stats': {
            'halls': len(dist),
            'assigned': sum(
                (1 if v['supervisor'] else 0) +
                sum(1 for m in v['monitors'] if m)
                for v in dist.values()
            ),
            'remaining': len(remaining),
            'gender_ok': sum(
                1 for v in dist.values()
                if any(m and m['gender'] == 'male' for m in v['monitors'])
                and any(m and m['gender'] == 'female' for m in v['monitors'])
            )
        }
    })


@app.route('/api/export/docx', methods=['POST'])
def api_export_docx():
    data = request.json
    distribution_raw = data.get('distribution', {})
    remaining_raw = data.get('remaining', [])
    day = data.get('day', '')
    date = data.get('date', '')
    session = data.get('session', '')

    docx_bytes = generate_word_document(
        distribution_raw, remaining_raw, day, date, session
    )
    buf = io.BytesIO(docx_bytes)
    buf.seek(0)
    filename = f"توزيع_المراقبين_{day}_{date or datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(buf, as_attachment=True,
                     download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == '__main__':
    import sys
    load_data()
    print()
    print("=" * 55)
    print("  الخادم يعمل على: http://localhost:5000")
    print("  افتح المتصفح على العنوان أعلاه")
    print("=" * 55)
    app.run(debug=False, host='0.0.0.0', port=5000)
